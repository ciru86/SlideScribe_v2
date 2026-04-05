#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Genera sia PDF che DOCX a partire da:
- slides.csv
- slide_texts.json
- immagini slide

Output:
- PDF orizzontale: Slide -> testo -> Slide -> testo ...
- DOCX orizzontale: Slide -> testo -> page break -> ...

Dipendenze:
    pip install pandas pillow reportlab python-docx
"""

import json
import re
import sys
import argparse
import logging
from pathlib import Path
from typing import Any

import pandas as pd
from PIL import Image
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.utils import ImageReader
from reportlab.pdfbase.pdfmetrics import stringWidth
from reportlab.pdfgen import canvas

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_BREAK




# =========================================================
# LOGGING
# =========================================================

logger = logging.getLogger("slides_and_texts_to_pdf")


def setup_logging(verbose: bool = False):
    # Configura un logging semplice su schermo.
    #
    # - INFO: mostra le fasi principali del lavoro
    # - DEBUG: mostra anche dettagli più fini utili per capire dove si trova lo script
    level = logging.DEBUG if verbose else logging.INFO

    # In molte esecuzioni questo script passa dentro wrapper shell con tee o pipe.
    # Usare stderr per i log rende la visualizzazione live più affidabile, mentre
    # stdout resta libero per eventuali messaggi finali o output "funzionale".
    if hasattr(sys.stderr, "reconfigure"):
        try:
            sys.stderr.reconfigure(line_buffering=True, write_through=True)
        except Exception:
            # Alcuni stream sostitutivi non supportano reconfigure: in quel caso
            # lasciamo il comportamento di default senza interrompere lo script.
            pass

    handler = logging.StreamHandler(sys.stderr)
    handler.setLevel(level)
    handler.setFormatter(logging.Formatter("[%(levelname)s] %(message)s"))

    logger.setLevel(level)
    logger.handlers.clear()
    logger.addHandler(handler)
    logger.propagate = False


# =========================================================
# UTILITY TEMPI / TESTO
# =========================================================

def seconds_to_hms(seconds: float) -> str:
    # Converte secondi float in stringa leggibile HH:MM:SS.
    #
    # Esempio:
    #   83.7 -> "00:01:23"
    #
    # Qui i millisecondi vengono ignorati volutamente perché:
    # - nel sommario non servono quasi mai
    # - il formato più corto è più leggibile
    total = int(seconds)
    h = total // 3600
    m = (total % 3600) // 60
    s = total % 60
    return f"{h:02d}:{m:02d}:{s:02d}"


def normalize_whitespace(text: str) -> str:
    # Normalizza spazi multipli, tab, newline ecc. in singoli spazi.
    #
    # Serve quando vuoi una riga "pulita" senza rumore tipografico.
    # Esempio:
    #   "ciao   \n   mondo" -> "ciao mondo"
    return re.sub(r"\s+", " ", text).strip()


def clean_final_text(text: str) -> str:
    """
    Pulizia leggera e conservativa del testo finale già corretto da LLM.
    Non fa magie e non prova a riscrivere nulla.
    """
    # Gestisce il caso None in ingresso.
    # Meglio convertirlo subito a stringa vuota.
    if text is None:
        return ""

    # Rimuove eventuale BOM e normalizza i newline.
    #
    # Questo evita:
    # - caratteri invisibili all'inizio del testo
    # - differenze Windows / Unix / vecchio Mac
    text = text.replace("\ufeff", "")
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = text.strip()

    # Pulisce i paragrafi uno per uno.
    #
    # Scelta voluta:
    # - comprime gli spazi dentro il paragrafo
    # - conserva la separazione in paragrafi se il testo la contiene davvero
    paragraphs = []

    for p in text.split("\n"):
        p = normalize_whitespace(p)
        if p:
            # Piccola rifinitura tipografica:
            # - toglie spazio prima della punteggiatura
            # - toglie spazi inutili subito dentro parentesi
            p = re.sub(r"\s+([,;:.!?])", r"\1", p)
            p = re.sub(r"\(\s+", "(", p)
            p = re.sub(r"\s+\)", ")", p)
            paragraphs.append(p)

    return "\n".join(paragraphs).strip()


def strip_markdown_inline(text: str) -> str:
    # Riduce il markdown inline a testo semplice leggibile.
    if not text:
        return ""

    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"__(.*?)__", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    text = re.sub(r"_(.*?)_", r"\1", text)
    return clean_final_text(text)


def load_summary_markdown(md_path: Path | None) -> list[dict[str, Any]]:
    # Legge un summary markdown semplice e lo trasforma in blocchi renderizzabili.
    if md_path is None:
        return []

    logger.info(f"Leggo il summary Markdown: {md_path}")
    raw = md_path.read_text(encoding="utf-8")
    raw = raw.replace("\ufeff", "").replace("\r\n", "\n").replace("\r", "\n").strip()

    if not raw:
        return []

    blocks: list[dict[str, Any]] = []
    paragraph_lines: list[str] = []

    def flush_paragraph() -> None:
        if not paragraph_lines:
            return
        text = strip_markdown_inline(" ".join(paragraph_lines))
        if text:
            blocks.append({"type": "paragraph", "text": text})
        paragraph_lines.clear()

    for raw_line in raw.split("\n"):
        line = raw_line.strip()

        if not line:
            flush_paragraph()
            continue

        heading_match = re.match(r"^(#{1,3})\s+(.*)$", line)
        if heading_match:
            flush_paragraph()
            level = min(len(heading_match.group(1)), 3)
            text = strip_markdown_inline(heading_match.group(2))
            if text:
                blocks.append({"type": "heading", "level": level, "text": text})
            continue

        list_match = re.match(r"^[-*]\s+(.*)$", line)
        if list_match:
            flush_paragraph()
            text = strip_markdown_inline(list_match.group(1))
            if text:
                blocks.append({"type": "list_item", "text": text})
            continue

        paragraph_lines.append(line)

    flush_paragraph()
    logger.info(f"Summary Markdown parsato: {len(blocks)} blocchi")
    return blocks


# =========================================================
# LETTURA INPUT STRUTTURATI
# =========================================================

def load_slide_texts_json(json_path: Path) -> dict[int, str]:
    # Legge il file JSON finale con i testi slide.
    #
    # Struttura attesa:
    # {
    #   "slides": [
    #       {"slide_index": 1, "text": "..."},
    #       {"slide_index": 2, "text": "..."}
    #   ]
    # }
    #
    # Restituisce:
    #   dizionario {slide_index: testo}
    logger.info(f"Leggo il JSON dei testi slide: {json_path}")
    raw = json.loads(json_path.read_text(encoding="utf-8"))

    # Validazione minima della root JSON.
    if not isinstance(raw, dict):
        raise ValueError(f"JSON non valido: root non è un oggetto -> {json_path}")

    slides = raw.get("slides")
    if not isinstance(slides, list):
        raise ValueError(f"JSON non valido: chiave 'slides' mancante o non lista -> {json_path}")

    slide_map: dict[int, str] = {}

    logger.info(f"Trovate {len(slides)} slide nel JSON dei testi")

    for item in slides:
        # Ogni elemento deve essere un dizionario.
        if not isinstance(item, dict):
            raise ValueError(f"Elemento slide non valido nel JSON: {item!r}")

        # Deve esserci sempre slide_index.
        if "slide_index" not in item:
            raise ValueError(f"Elemento slide senza 'slide_index': {item!r}")

        slide_index = int(item["slide_index"])
        text = item.get("text", "")

        # Se il testo manca o è None, lo trattiamo come stringa vuota.
        if text is None:
            text = ""

        # Pulizia finale del testo prima dell'uso.
        text = clean_final_text(str(text))

        # Difesa contro duplicati nel JSON.
        if slide_index in slide_map:
            raise ValueError(f"Slide duplicata nel JSON: {slide_index}")

        slide_map[slide_index] = text

    logger.info(f"Mappa testi costruita: {len(slide_map)} slide indicizzate")
    return slide_map


def build_entries_from_csv_and_json(slides_df: pd.DataFrame, slide_text_map: dict[int, str]):
    # Costruisce la struttura dati unificata che useranno sia PDF sia DOCX.
    #
    # Input:
    # - slides_df: dataframe letto da slides.csv
    # - slide_text_map: dizionario {slide_index: testo}
    #
    # Output:
    # - lista di dict, uno per slide, con:
    #   slide_index, slide_start, slide_end, filename, text
    #
    # slide_end viene derivato dal timestamp della slide successiva.
    # Per l'ultima slide si usa +inf, perché non c'è una successiva.
    logger.info("Costruisco la struttura unificata slide/testi")
    slides = slides_df.copy().sort_values("timestamp_sec").reset_index(drop=True)

    required_cols = {"slide_index", "timestamp_sec", "filename"}
    missing = required_cols - set(slides.columns)
    if missing:
        raise ValueError(f"Nel CSV mancano le colonne richieste: {sorted(missing)}")

    entries = []

    timestamps = slides["timestamp_sec"].tolist()
    filenames = slides["filename"].tolist()
    slide_indices = slides["slide_index"].tolist()

    logger.debug(f"CSV ordinato: {len(slides)} righe")

    for i in range(len(slides)):
        slide_index = int(slide_indices[i])
        start_t = float(timestamps[i])
        end_t = float(timestamps[i + 1]) if i < len(slides) - 1 else float("inf")
        filename = str(filenames[i])

        # Se una slide non ha testo nel JSON, qui finisce come stringa vuota.
        text = slide_text_map.get(slide_index, "")

        entries.append({
            "slide_index": slide_index,
            "slide_start": start_t,
            "slide_end": end_t,
            "filename": filename,
            "text": text,
        })

    logger.info(f"Struttura finale pronta: {len(entries)} slide")
    return entries


# =========================================================
# WRAPPING TESTO PDF
# =========================================================

def wrap_text_to_width(text, font_name, font_size, max_width):
    # Esegue il line wrapping manuale del testo per il PDF.
    #
    # reportlab, in questo flusso, non sta usando un motore paragrafi avanzato.
    # Quindi qui spezziamo noi il testo in righe che stiano dentro max_width.
    #
    # Logica:
    # 1) divide per paragrafi
    # 2) aggiunge parole finché la riga ci sta
    # 3) quando sfora, va a capo
    # 4) se una singola parola è troppo lunga, la spezza a caratteri
    if not text:
        return []

    paragraphs = [p.strip() for p in text.split("\n") if p.strip()]
    all_lines = []

    for p in paragraphs:
        words = p.split()
        current = ""

        for word in words:
            test = word if not current else current + " " + word
            width = stringWidth(test, font_name, font_size)

            # Se la riga testata ci sta, la accumuliamo.
            if width <= max_width:
                current = test
            else:
                # La riga corrente è piena: la salviamo.
                if current:
                    all_lines.append(current)

                # Se la parola da sola ci sta, parte una nuova riga.
                if stringWidth(word, font_name, font_size) <= max_width:
                    current = word
                else:
                    # Caso raro ma possibile:
                    # una singola parola è più larga della colonna.
                    #
                    # Allora la spezziamo a caratteri.
                    chunk = ""
                    for ch in word:
                        test_chunk = chunk + ch
                        if stringWidth(test_chunk, font_name, font_size) <= max_width:
                            chunk = test_chunk
                        else:
                            if chunk:
                                all_lines.append(chunk)
                            chunk = ch
                    current = chunk

        if current:
            all_lines.append(current)

        # Riga vuota tra paragrafi.
        all_lines.append("")

    # Rimuove eventuali righe vuote finali inutili.
    while all_lines and all_lines[-1] == "":
        all_lines.pop()

    return all_lines


# =========================================================
# PDF HELPERS
# =========================================================

def fit_image_in_box(img_w, img_h, box_w, box_h):
    # Calcola la dimensione massima dell'immagine mantenendo il rapporto d'aspetto.
    #
    # Restituisce:
    #   (draw_w, draw_h)
    #
    # Viene usato per far stare la slide dentro il box assegnato senza deformarla.
    scale = min(box_w / img_w, box_h / img_h)
    return img_w * scale, img_h * scale


def draw_footer(c, page_w, page_h, page_num):
    # Disegna il numero pagina nel footer in basso a destra.
    #
    # page_h qui non viene usato, ma tenerlo in firma rende la funzione
    # coerente con le altre helper di disegno.
    c.setFont("Helvetica", 9)
    c.drawRightString(page_w - 24, 16, str(page_num))


def draw_header(c, page_w, page_h, title_left, title_right=None):
    # Disegna intestazione standard della pagina PDF:
    # - titolo a sinistra
    # - eventuale info a destra
    # - linea orizzontale di separazione sotto l'header
    c.setFont("Helvetica-Bold", 12)
    c.drawString(24, page_h - 20, title_left)

    if title_right:
        c.setFont("Helvetica", 10)
        c.drawRightString(page_w - 24, page_h - 20, title_right)

    c.line(24, page_h - 26, page_w - 24, page_h - 26)


def draw_cover_page(c, page_w, page_h, pdf_title, total_slides, source_url, page_num):
    # Disegna una cover iniziale con titolo, numero slide e URL sorgente.
    logger.info("Genero la cover PDF")
    draw_footer(c, page_w, page_h, page_num)

    center_x = page_w / 2
    y = page_h - 90

    c.setFont("Helvetica-Bold", 24)
    c.drawCentredString(center_x, y, pdf_title)
    y -= 34

    c.setFont("Helvetica", 13)
    c.drawCentredString(center_x, y, f"Numero slide: {total_slides}")
    y -= 42

    if source_url:
        c.setFont("Helvetica-Bold", 13)
        c.drawCentredString(center_x, y, "Sorgente video")
        y -= 24

        c.setLineWidth(1)
        c.setStrokeColorRGB(0.30, 0.42, 0.58)
        c.roundRect(60, y - 26, page_w - 120, 44, 8, stroke=1, fill=0)
        c.setStrokeColorRGB(0, 0, 0)

        c.setFont("Helvetica", 10)
        url_lines = wrap_text_to_width(source_url, "Helvetica", 10, page_w - 150)
        line_y = y
        for line in url_lines[:2]:
            c.drawCentredString(center_x, line_y, line)
            line_y -= 13

    c.showPage()
    return page_num + 1


def draw_summary_pages(c, page_w, page_h, entries, pdf_title, start_page_num):
    """
    Disegna il sommario su una o più pagine PDF.
    Restituisce il prossimo page number disponibile.
    """
    logger.info(f"Genero il sommario PDF ({len(entries)} slide)")
    page_num = start_page_num

    # Margini e metrica verticale.
    #
    # top_y in realtà qui non viene usato dopo l'assegnazione,
    # ma tenerlo leggibile aiuta a capire l'impostazione pagina.
    margin_left = 36
    margin_right = 36
    top_y = page_h - 40
    bottom_y = 32
    line_h = 12

    title_drawn = False
    i = 0
    total = len(entries)

    while i < total:
        draw_header(c, page_w, page_h, pdf_title, "Sommario")
        draw_footer(c, page_w, page_h, page_num)

        y = page_h - 70

        # Prima pagina del sommario: titolo grande + numero slide.
        if not title_drawn:
            c.setFont("Helvetica-Bold", 20)
            c.drawString(margin_left, y, pdf_title)
            y -= 25

            c.setFont("Helvetica", 12)
            c.drawString(margin_left, y, f"Numero slide: {len(entries)}")
            y -= 40

            title_drawn = True
        else:
            # Pagine successive del sommario.
            c.setFont("Helvetica-Bold", 14)
            c.drawString(margin_left, y, "Indice (continuazione)")
            y -= 30

        c.setFont("Helvetica-Bold", 12)
        c.drawString(margin_left, y, "Indice")
        y -= 22

        c.setFont("Helvetica", 9)

        # Scrive le righe del sommario finché c'è spazio utile verticale.
        while i < total and y >= bottom_y + line_h:
            entry = entries[i]
            row = (
                f"Slide {entry['slide_index']:>3}   "
                f"[{seconds_to_hms(entry['slide_start'])}]   "
                f"{entry['filename']}"
            )
            c.drawString(margin_left + 12, y, row)
            y -= line_h
            i += 1

        c.showPage()
        page_num += 1

    return page_num


def build_summary_render_blocks(summary_blocks, page_w):
    # Converte i blocchi markdown in righe con metriche adatte al PDF.
    render_blocks = []
    max_width = page_w - 72
    current_section = ""

    for block in summary_blocks:
        block_type = block["type"]

        if block_type == "heading":
            level = int(block.get("level", 2))
            current_section = block["text"].strip().lower()
            if level <= 1:
                font_name = "Helvetica-Bold"
                font_size = 20
                line_height = 26
                spacing_before = 10
                spacing_after = 12
            elif level == 2:
                font_name = "Helvetica-Bold"
                font_size = 15
                line_height = 20
                spacing_before = 12
                spacing_after = 6
            else:
                font_name = "Helvetica-Bold"
                font_size = 12
                line_height = 16
                spacing_before = 6
                spacing_after = 3

            lines = wrap_text_to_width(block["text"], font_name, font_size, max_width)
            render_blocks.append({
                "type": "heading",
                "lines": lines,
                "font_name": font_name,
                "font_size": font_size,
                "line_height": line_height,
                "spacing_before": spacing_before,
                "spacing_after": spacing_after,
                "indent": 0,
                "accent": "implicazioni cliniche" in current_section,
            })
            continue

        if block_type == "list_item":
            text = f"- {block['text']}"
            font_name = "Helvetica"
            font_size = 11
            line_height = 15
            lines = wrap_text_to_width(text, font_name, font_size, max_width - 12)
            render_blocks.append({
                "type": "list_item",
                "lines": lines,
                "font_name": font_name,
                "font_size": font_size,
                "line_height": line_height,
                "spacing_before": 3,
                "spacing_after": 3,
                "indent": 16,
                "accent": False,
            })
            continue

        font_name = "Helvetica"
        font_size = 11
        line_height = 16
        is_clinical = "implicazioni cliniche" in current_section
        indent = 12 if is_clinical else 0
        width = max_width - 18 if is_clinical else max_width
        lines = wrap_text_to_width(block["text"], font_name, font_size, width)
        render_blocks.append({
            "type": "paragraph",
            "lines": lines,
            "font_name": font_name,
            "font_size": font_size,
            "line_height": line_height,
            "spacing_before": 3,
            "spacing_after": 8,
            "indent": indent,
            "accent": is_clinical,
        })

    return render_blocks


def draw_markdown_summary_pages(c, page_w, page_h, summary_blocks, pdf_title, start_page_num):
    # Disegna il riassunto markdown su una o più pagine PDF.
    if not summary_blocks:
        return start_page_num

    logger.info(f"Genero il summary PDF ({len(summary_blocks)} blocchi)")
    render_blocks = build_summary_render_blocks(summary_blocks, page_w)

    margin_left = 36
    bottom_y = 32
    page_num = start_page_num
    i = 0

    while i < len(render_blocks):
        draw_header(c, page_w, page_h, pdf_title, "Riassunto")
        draw_footer(c, page_w, page_h, page_num)

        y = page_h - 60

        while i < len(render_blocks):
            block = render_blocks[i]
            needed_h = block["spacing_before"] + len(block["lines"]) * block["line_height"] + block["spacing_after"]

            if y - needed_h < bottom_y:
                break

            y -= block["spacing_before"]

            if block.get("accent"):
                line_top = y + 3
                line_bottom = y - len(block["lines"]) * block["line_height"] - block["spacing_after"] + 4
                c.setLineWidth(2)
                c.setStrokeColorRGB(0.30, 0.42, 0.58)
                c.line(margin_left - 8, line_top, margin_left - 8, line_bottom)
                c.setLineWidth(1)
                c.setStrokeColorRGB(0, 0, 0)

            c.setFont(block["font_name"], block["font_size"])

            for line in block["lines"]:
                c.drawString(margin_left + block["indent"], y, line)
                y -= block["line_height"]

            y -= block["spacing_after"]
            i += 1

        c.showPage()
        page_num += 1

    return page_num


def draw_slide_page(c, page_w, page_h, slide_path, text_lines, slide_num, time_range, page_num):
    # Disegna la pagina principale di una slide nel PDF:
    # - header
    # - footer
    # - immagine della slide
    # - testo sotto l'immagine
    #
    # Restituisce il numero di righe testo effettivamente consumate
    # in questa pagina.
    logger.debug(f"Disegno pagina PDF della slide {slide_num} ({slide_path.name})")
    draw_header(c, page_w, page_h, f"Slide {slide_num}", time_range)
    draw_footer(c, page_w, page_h, page_num)

    margin_x = 24
    top_y = page_h - 40
    bottom_y = 32

    usable_w = page_w - 2 * margin_x
    usable_h = top_y - bottom_y

    # Il 70% dell'altezza utile viene riservato all'immagine della slide.
    # Il resto va al testo.
    img_area_h = usable_h * 0.70
    gap = 10

    img_box_w = usable_w
    img_box_h = img_area_h

    if slide_path.exists():
        # Legge le dimensioni dell'immagine solo per calcolare il fit.
        with Image.open(slide_path) as im:
            img_w, img_h = im.size

        draw_w, draw_h = fit_image_in_box(img_w, img_h, img_box_w, img_box_h)

        # Centra orizzontalmente l'immagine nel box disponibile.
        x_img = margin_x + (img_box_w - draw_w) / 2
        y_img = top_y - draw_h

        c.drawImage(
            ImageReader(str(slide_path)),
            x_img,
            y_img,
            width=draw_w,
            height=draw_h,
            preserveAspectRatio=True,
            mask='auto'
        )
    else:
        # Fallback se l'immagine non esiste.
        c.setFont("Helvetica-Bold", 16)
        c.drawString(margin_x, top_y - 30, f"[Immagine non trovata: {slide_path.name}]")
        y_img = top_y - 60

    font_name = "Helvetica"
    font_size = 11
    line_height = 14
    c.setFont(font_name, font_size)

    # Il testo parte subito sotto l'immagine.
    text_top_y = y_img - gap
    available_h = text_top_y - bottom_y
    max_lines = max(1, int(available_h // line_height))

    consumed = min(len(text_lines), max_lines)

    y = text_top_y - line_height + 2
    for line in text_lines[:consumed]:
        c.drawString(margin_x, y, line)
        y -= line_height

    return consumed


def draw_text_continuation_page(c, page_w, page_h, text_lines, slide_num, page_num):
    # Disegna una pagina aggiuntiva di solo testo, usata quando il testo
    # della slide non entra tutto nella pagina principale.
    #
    # Restituisce quante righe sono state consumate in questa pagina.
    logger.debug(f"Aggiungo pagina PDF di continuazione testo per slide {slide_num}")
    draw_header(c, page_w, page_h, f"Slide {slide_num}", "Continuazione testo")
    draw_footer(c, page_w, page_h, page_num)

    margin_x = 36
    top_y = page_h - 40
    bottom_y = 32

    font_name = "Helvetica"
    font_size = 11
    line_height = 14
    c.setFont(font_name, font_size)

    available_h = top_y - bottom_y
    max_lines = max(1, int(available_h // line_height))

    consumed = min(len(text_lines), max_lines)

    y = top_y - line_height
    for line in text_lines[:consumed]:
        c.drawString(margin_x, y, line)
        y -= line_height

    return consumed


def build_pdf(entries, input_dir: Path, output_pdf: Path, summary_blocks=None, source_url: str = ""):
    # Costruisce l'intero PDF finale.
    #
    # Flusso:
    # 1) crea canvas A4 landscape
    # 2) genera pagine di sommario
    # 3) per ogni slide crea almeno una pagina
    # 4) se il testo è troppo lungo, aggiunge pagine di continuazione
    # 5) salva il PDF
    logger.info(f"Avvio generazione PDF: {output_pdf}")
    page_w, page_h = landscape(A4)
    c = canvas.Canvas(str(output_pdf), pagesize=(page_w, page_h))

    # Titolo "umano" derivato dal nome file output.
    pdf_title = output_pdf.stem.replace("_", " ").replace("-", " ")

    page_num = 1
    page_num = draw_cover_page(c, page_w, page_h, pdf_title, len(entries), source_url, page_num)
    page_num = draw_markdown_summary_pages(c, page_w, page_h, summary_blocks or [], pdf_title, page_num)

    logger.info(f"Genero le pagine DOCX delle slide ({len(entries)} totali)")

    for entry in entries:
        slide_path = input_dir / entry["filename"]
        raw_text = entry["text"].strip()

        # Placeholder se il testo è vuoto.
        if not raw_text:
            raw_text = "[Nessun testo associato a questa slide]"

        max_text_width = page_w - 72
        lines = wrap_text_to_width(raw_text, "Helvetica", 11, max_text_width)

        time_range = f"{seconds_to_hms(entry['slide_start'])}"

        logger.info(f"PDF -> slide {entry['slide_index']} | immagine: {slide_path.name}")

        consumed = draw_slide_page(
            c,
            page_w,
            page_h,
            slide_path=slide_path,
            text_lines=lines,
            slide_num=entry["slide_index"],
            time_range=time_range,
            page_num=page_num
        )
        c.showPage()
        page_num += 1

        # Se resta testo non ancora disegnato, lo spalma su pagine successive.
        remaining = lines[consumed:]

        while remaining:
            logger.info(f"PDF -> slide {entry['slide_index']} richiede una pagina di continuazione")
            consumed2 = draw_text_continuation_page(
                c,
                page_w,
                page_h,
                text_lines=remaining,
                slide_num=entry["slide_index"],
                page_num=page_num
            )
            c.showPage()
            page_num += 1
            remaining = remaining[consumed2:]

    page_num = draw_summary_pages(c, page_w, page_h, entries, pdf_title, page_num)

    c.save()
    logger.info(f"PDF salvato: {output_pdf}")


# =========================================================
# DOCX HELPERS
# =========================================================

def set_landscape(document: Document):
    # Imposta il documento Word in orizzontale e regola i margini.
    #
    # Nota:
    # in python-docx, per cambiare davvero orientamento bisogna anche
    # scambiare page_width e page_height.
    logger.debug("Imposto il DOCX in formato landscape")
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width

    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)


def get_usable_width_inches(document: Document) -> float:
    # Calcola la larghezza utile effettiva della pagina, al netto dei margini.
    #
    # python-docx lavora internamente in EMU.
    # Qui convertiamo in pollici perché add_picture usa Inches(...).
    logger.debug("Imposto il DOCX in formato landscape")
    section = document.sections[0]
    usable = section.page_width - section.left_margin - section.right_margin
    return usable / 914400  # EMU -> inches


def add_docx_summary(document: Document, entries, title: str):
    # Aggiunge al DOCX una sezione iniziale di sommario:
    # - titolo
    # - numero slide
    # - elenco slide con timestamp e filename
    # - page break finale
    logger.info(f"Aggiungo sommario DOCX ({len(entries)} slide)")

    p = document.add_paragraph()
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(18)

    p = document.add_paragraph(f"Numero slide: {len(entries)}")
    p.runs[0].font.size = Pt(11)

    p = document.add_paragraph()
    r = p.add_run("Indice")
    r.bold = True
    r.font.size = Pt(13)

    for entry in entries:
        p = document.add_paragraph(
            f"Slide {entry['slide_index']}  [{seconds_to_hms(entry['slide_start'])}]  {entry['filename']}"
        )
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(10)

    document.add_page_break()


def add_docx_cover(document: Document, title: str, total_slides: int, source_url: str):
    # Aggiunge una cover iniziale con titolo e sorgente video.
    logger.info("Aggiungo cover DOCX")

    p = document.add_paragraph()
    p.alignment = 1
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(20)

    p = document.add_paragraph()
    p.alignment = 1
    r = p.add_run(f"Numero slide: {total_slides}")
    r.font.size = Pt(12)

    if source_url:
        p = document.add_paragraph()
        p.alignment = 1
        r = p.add_run("Sorgente video")
        r.bold = True
        r.font.size = Pt(12)

        p = document.add_paragraph()
        p.alignment = 1
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.right_indent = Inches(0.5)
        r = p.add_run(source_url)
        r.font.size = Pt(10)

    document.add_page_break()


def add_docx_markdown_summary(document: Document, summary_blocks):
    # Aggiunge il riassunto markdown al DOCX con una resa semplice e leggibile.
    if not summary_blocks:
        return

    logger.info(f"Aggiungo summary DOCX ({len(summary_blocks)} blocchi)")
    current_section = ""

    for block in summary_blocks:
        block_type = block["type"]
        text = block["text"]

        if block_type == "heading":
            level = int(block.get("level", 2))
            current_section = text.strip().lower()
            p = document.add_paragraph()
            r = p.add_run(text)
            r.bold = True
            if level <= 1:
                r.font.size = Pt(18)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(10)
            elif level == 2:
                r.font.size = Pt(14)
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(5)
            else:
                r.font.size = Pt(11)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(3)
            if "implicazioni cliniche" in current_section:
                p.paragraph_format.left_indent = Inches(0.12)
            continue

        if block_type == "list_item":
            try:
                p = document.add_paragraph(style="List Bullet")
            except KeyError:
                p = document.add_paragraph()
                p.add_run("- ")
            p.add_run(text)
            p.paragraph_format.space_after = Pt(3)
            for run in p.runs:
                run.font.size = Pt(11)
            continue

        p = document.add_paragraph(text)
        p.paragraph_format.space_after = Pt(8)
        if "implicazioni cliniche" in current_section:
            p.paragraph_format.left_indent = Inches(0.18)
        for run in p.runs:
            run.font.size = Pt(11)

    document.add_page_break()


def add_slide_block_docx(document: Document, slide_num: int, slide_path: Path, text: str, max_image_width_inches: float):
    # Aggiunge al DOCX un blocco completo per una slide:
    # - titolo slide
    # - immagine slide
    # - testo associato
    # - page break finale
    logger.info(f"DOCX -> slide {slide_num} | immagine: {slide_path.name}")
    p = document.add_paragraph()
    r = p.add_run(f"Slide {slide_num}")
    r.bold = True
    r.font.size = Pt(14)

    if slide_path.exists():
        try:
            # Ridimensiona l'immagine in larghezza.
            # L'altezza viene adattata automaticamente mantenendo le proporzioni.
            document.add_picture(str(slide_path), width=Inches(max_image_width_inches))
        except Exception:
            # Fallback se l'immagine esiste ma Word/python-docx non riesce a usarla.
            p = document.add_paragraph(f"[Errore nel caricamento immagine: {slide_path.name}]")
            for run in p.runs:
                run.font.size = Pt(10)
    else:
        # Fallback se il file immagine manca del tutto.
        p = document.add_paragraph(f"[Immagine non trovata: {slide_path.name}]")
        for run in p.runs:
            run.font.size = Pt(10)

    p = document.add_paragraph(text if text else "[Nessun testo associato a questa slide]")
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(0)
    for run in p.runs:
        run.font.size = Pt(11)

    # Page break: una slide per pagina/blocco.
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def build_docx(entries, input_dir: Path, output_docx: Path, summary_blocks=None, source_url: str = ""):
    # Costruisce il DOCX finale.
    #
    # Flusso:
    # 1) crea documento
    # 2) imposta landscape
    # 3) aggiunge sommario
    # 4) aggiunge un blocco per ogni slide
    # 5) salva
    logger.info(f"Avvio generazione DOCX: {output_docx}")
    doc = Document()
    set_landscape(doc)

    title = output_docx.stem.replace("_", " ").replace("-", " ")
    add_docx_cover(doc, title, len(entries), source_url)
    add_docx_markdown_summary(doc, summary_blocks or [])

    # Larghezza massima immagine leggermente inferiore alla larghezza utile
    # per evitare che l'immagine tocchi troppo i bordi.
    usable_width = get_usable_width_inches(doc)
    image_width = max(4.0, usable_width - 0.2)

    logger.info(f"Genero le pagine PDF delle slide ({len(entries)} totali)")

    for entry in entries:
        slide_path = input_dir / entry["filename"]
        raw_text = entry["text"].strip()

        add_slide_block_docx(
            doc,
            slide_num=entry["slide_index"],
            slide_path=slide_path,
            text=raw_text,
            max_image_width_inches=image_width
        )

    add_docx_summary(doc, entries, title)

    doc.save(str(output_docx))
    logger.info(f"DOCX salvato: {output_docx}")


# =========================================================
# MAIN
# =========================================================

def main():
    # Parser CLI principale dello script.
    #
    # Questo script si aspetta:
    # - una cartella input con slides.csv e immagini
    # - un JSON con i testi finali slide
    # - un nome base per gli output
    parser = argparse.ArgumentParser(
        description="Unisce slides.csv + slide_texts.json + immagini slide in PDF e DOCX."
    )

    parser.add_argument(
        "--input-dir",
        default=".",
        help="Cartella contenente slides.csv e immagini slide",
    )
    parser.add_argument(
        "--csv",
        default="slides.csv",
        help="Nome file CSV delle slide",
    )
    parser.add_argument(
        "--slide-texts",
        required=True,
        help="Path del file JSON con il testo finale per slide",
    )
    parser.add_argument(
        "--summary-file",
        help="Path opzionale del file Markdown con il riassunto finale",
    )
    parser.add_argument(
        "--youtube-url",
        help="URL YouTube opzionale da mostrare nella cover iniziale",
    )
    parser.add_argument(
        "--output-base",
        default="slides_con_testo",
        help="Nome base output senza estensione",
    )
    parser.add_argument(
        "-v", "--verbose",
        action="store_true",
        help="Mostra log più dettagliati su schermo",
    )
    args = parser.parse_args()

    setup_logging(args.verbose)
    logger.info("Avvio script slides_and_texts_to_pdf")

    # Risolve i path in modo robusto.
    input_dir = Path(args.input_dir.strip()).expanduser().resolve()
    logger.info(f"Cartella input: {input_dir}")
    csv_path = input_dir / args.csv
    slide_texts_path = Path(args.slide_texts).expanduser().resolve()
    summary_path = Path(args.summary_file).expanduser().resolve() if args.summary_file else None
    youtube_url = args.youtube_url.strip() if args.youtube_url else ""

    output_pdf = input_dir / f"{args.output_base}.pdf"
    output_docx = input_dir / f"{args.output_base}.docx"

    logger.info(f"CSV atteso: {csv_path}")
    logger.info(f"JSON testi: {slide_texts_path}")
    logger.info(f"Summary:     {summary_path if summary_path else '[assente]'}")
    logger.info(f"YouTube URL:  {youtube_url if youtube_url else '[assente]'}")
    logger.info(f"Output PDF:  {output_pdf}")
    logger.info(f"Output DOCX: {output_docx}")

    # Validazione esistenza input.
    if not input_dir.exists():
        raise FileNotFoundError(f"Cartella non trovata: {input_dir}")

    if not csv_path.exists():
        raise FileNotFoundError(f"CSV non trovato: {csv_path}")

    if not slide_texts_path.exists():
        raise FileNotFoundError(f"JSON slide texts non trovato: {slide_texts_path}")

    if summary_path is not None and not summary_path.exists():
        raise FileNotFoundError(f"Summary file non trovato: {summary_path}")

    # Lettura CSV.
    logger.info(f"Leggo il CSV slide: {csv_path}")
    slides_df = pd.read_csv(csv_path)
    logger.info(f"CSV caricato: {len(slides_df)} righe")

    required_cols = {"slide_index", "timestamp_sec", "filename"}
    missing = required_cols - set(slides_df.columns)
    if missing:
        raise ValueError(f"Nel CSV mancano le colonne richieste: {sorted(missing)}")

    # Lettura testi e costruzione struttura unificata.
    slide_text_map = load_slide_texts_json(slide_texts_path)
    entries = build_entries_from_csv_and_json(slides_df, slide_text_map)
    summary_blocks = load_summary_markdown(summary_path)

    # Generazione output finali.
    logger.info("Inizio generazione artefatti finali")
    build_pdf(entries, input_dir, output_pdf, summary_blocks=summary_blocks, source_url=youtube_url)
    build_docx(entries, input_dir, output_docx, summary_blocks=summary_blocks, source_url=youtube_url)

    logger.info("Elaborazione completata")
    print(f"PDF creato:  {output_pdf}")
    print(f"DOCX creato: {output_docx}")


if __name__ == "__main__":
    main()
